from flask import Flask, request, jsonify
from flask_cors import CORS
import os
from dotenv import load_dotenv
from pathlib import Path
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.http import models
import google.generativeai as genai
from langchain.document_loaders import PyPDFLoader, TextLoader
from langchain.schema import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings

from llm_response import get_response , ask_mor


# Flask App Configuration
app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

load_dotenv()
# Fetch API key from environment variables
GENAI_API_KEY = os.getenv("GENAI_API_KEY")

if not GENAI_API_KEY:
    raise ValueError("GENAI_API_KEY is not set in the .env file")
genai.configure(api_key=GENAI_API_KEY)

generation_config_scene = {
    "temperature": 0.9,
    "top_p": 1,
    "top_k": 1,
    "max_output_tokens": 300,
}

generation_config_text = {
    "temperature": 0.7,
    "top_p": 1,
    "top_k": 1,
    "max_output_tokens": 500,
}

# Gemini Model Initialization
scene_model = genai.GenerativeModel(
    "gemini-1.5-flash",
    generation_config=generation_config_scene
)

text_model = genai.GenerativeModel(
    "gemini-1.5-flash",
    generation_config=generation_config_text
)

# Embedding Model Initialization
embedding = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
vectorstore = None  # FAISS Vector Store initialized dynamically
retriever = None

# ----------------- Helper Functions -----------------

def load_document(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
        return loader.load()
    elif ext == ".txt":
        loader = TextLoader(file_path)
        return loader.load()
    elif ext == ".docx":
        docx_doc = DocxDocument(file_path)
        docs = [
            LangchainDocument(page_content=para.text.strip(), metadata={"source": file_path})
            for para in docx_doc.paragraphs if para.text.strip()
        ]
        return docs
    else:
        raise ValueError("Unsupported file type.")

def process_and_store_document(file_path):
    global vectorstore, retriever
    documents = load_document(file_path)
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    chunks = text_splitter.split_documents(documents)

    if vectorstore is None:
        vectorstore = FAISS.from_documents(chunks, embedding)
    else:
        vectorstore.add_documents(chunks)

    retriever = vectorstore.as_retriever()

# ----------------- API Routes from App 1 -----------------

@app.route('/api/scene-description', methods=['POST'])
def scene_description():
    if 'image' not in request.files:
        return jsonify({"error": "No image file uploaded"}), 400

    image = request.files['image']
    if not image:
        return jsonify({"error": "No file provided"}), 400

    image_path = Path(os.path.join(UPLOAD_FOLDER, image.filename))
    image.save(image_path)

    try:
        image_part = {
            "mime_type": image.content_type,
            "data": image_path.read_bytes()
        }
        prompt_parts = [
            "Describe this image as if narrating to a blind user for 15 seconds. Provide a short and concise scene description. Avoid any visual references like 'look at that' or 'you can see.' Instead, focus on auditory and descriptive language that conveys the context. If there is an interesting item, name it explicitly and add the following dialog: 'There is [interesting item] in this scene. Do you want to learn about it?' Additionally, provide a teaching-style explanation about the interesting item in the following format:\n\n",
            "Scene Description: <short scene description ending with the interesting item and a dialog>\n",
            "Learning: <teaching-style explanation about the interesting item>\n\n",
            image_part
        ]

        response = scene_model.generate_content(prompt_parts)

        scene_description = ""
        learning_content = ""
        for chunk in response:
            text = chunk.text
            if "Scene Description:" in text:
                scene_description = text.split("Scene Description:")[1].split("Learning:")[0].strip()
            if "Learning:" in text:
                learning_content = text.split("Learning:")[1].strip()

        has_learning = bool(learning_content)

        return jsonify({
            "scene_description": scene_description,
            "has_learning": has_learning,
            "learning": learning_content
        })
    finally:
        if image_path.exists():
            os.remove(image_path)

@app.route('/api/command', methods=['POST'])
def process_command():
    data = request.get_json()
    command = data.get("command", "").lower()

    if command == "yes":
        return jsonify({
            "status": "success",
            "message": "You chose to learn more!",
            "learning": "Please refer to the provided learning content."
        })
    elif command == "no":
        return jsonify({
            "status": "success",
            "message": "Acknowledgment received. No further action."
        })
    else:
        return jsonify({
            "status": "error",
            "message": "Invalid command."
        }), 400

# ----------------- API Routes from App 2 -----------------

@app.route('/upload', methods=['POST'])
def upload_document():
    if 'file' not in request.files:
        return jsonify({"status": "error", "message": "No file uploaded"}), 400

    file = request.files['file']
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(file_path)

    try:
        process_and_store_document(file_path)
        return jsonify({"status": "success", "message": "Document uploaded and processed successfully."})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500
    finally:
        os.remove(file_path)

@app.route('/query', methods=['POST'])
def handle_query():
    data = request.get_json()
    user_query = data.get("query", "")

    if not user_query:
        return jsonify({"status": "error", "message": "No query provided"}), 400

    try:
        prompt = f"Answer this query concisely:\n{user_query}"
        response = text_model.generate_content([prompt])
        generated_text = "".join([chunk.text for chunk in response])
        return jsonify({"status": "success", "response": generated_text})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)})

@app.route('/summarize', methods=['POST'])
def summarize_document():
    global retriever
    if retriever is None:
        return jsonify({"status": "error", "message": "No documents uploaded yet."}), 400

    try:
        results = retriever.get_relevant_documents("Summarize the document.")
        if not results:
            return jsonify({"status": "error", "message": "No relevant documents found."}), 400

        context = "\n".join([doc.page_content for doc in results])

        prompt = f"Context: {context}\n\nSummarize the content concisely:"
        response = text_model.generate_content([prompt])
        summary = "".join([chunk.text for chunk in response])

        return jsonify({"status": "success", "summary": summary})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/api/scrap_subject', methods=['POST'])
def scrap_sub():
    try:
        data = request.json
        subject = data.get('subject')
        print(subject)

        if not subject:
            return jsonify({'error': 'Subject not provided'}), 400

        description = get_response(subject)
        print(description)

        return jsonify({'success': True, 'description': description})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/ask_more', methods=['POST'])
def handle_ask_more():
    """
    Endpoint to handle conversation queries.
    """
    data = request.get_json()
    query = data.get('query')
    context = data.get('conversation_history')
    
    response = ask_mor(context,query)
    return jsonify(response)
# ----------------- Run Flask App -----------------
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
